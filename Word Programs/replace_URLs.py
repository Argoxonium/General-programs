import re
import os
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

def main() -> None:
    #Get the path and open the word document
    path = input('What is the path location for the word document')

    #open the word document from the inputted path
    doc = open_word_document(path)

    #extract all the hypelinks within the document
    hyperlinks = extract_hyperlinks(doc)

    #change the hyperlinks with the new goal code in mind
    new_hyperlinks = change_url(hyperlinks,"https://my.anl.gov/esb/view/")

    #update urls within the document
    new_doc = update_urls_in_doc(doc,new_hyperlinks)

    #create a new path
    save_path = new_path(path)
    
    #save the document
    save_doc(new_doc,path)

def open_word_document(file_path: str) -> Document:
    """
    Opens a Word document (.docx) from the specified file path and returns a Document object.

    :param file_path: The path to the Word document file to be opened.
    :type file_path: str
    :return: A Document object representing the Word document.
    :rtype: docx.Document

    :raises FileNotFoundError: If the file at the specified path does not exist.
    :raises ValueError: If the file path does not point to a .docx file.
    """
    # Check if the file path is valid and points to a .docx file
    if not file_path.endswith('.docx'):
        raise ValueError("The file path must point to a .docx file.")

    # Attempt to open the document
    try:
        doc = Document(file_path)
        print(f"Successfully opened the document: {file_path}")
        return doc
    except FileNotFoundError:
        print(f"Error: The file at {file_path} was not found.")
        raise
    except Exception as e:
        print(f"An error occurred while opening the document: {e}")
        raise

def extract_hyperlinks(doc: Document) -> dict:
    """
    Extracts all hyperlinks from a Word document and returns them in a dictionary with 
    the hyperlink text as keys and the URLs as values.

    :param doc: The document that was inputted in from the user.
    :type doc: Document
    :return: A dictionary containing hyperlink texts and their corresponding URLs.
    :rtype: dict
    """

    hyperlinks = {}  # Dictionary to store hyperlinks and their texts

    # Iterate through the document's part relationships
    rels = doc.part.rels  # Access the document's relationship dictionary
    
    # Iterate through each paragraph in the document
    for paragraph in doc.paragraphs:
        # Parse the paragraph XML
        p_xml = paragraph._element
        
        # Find all hyperlink tags within the paragraph
        for hyperlink in p_xml.findall(".//w:hyperlink", namespaces=p_xml.nsmap):
            # Extract the relationship ID of the hyperlink
            r_id = hyperlink.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
            
            if r_id and r_id in rels:
                # Extract the URL using the relationship ID
                url = rels[r_id]._target
                
                # Extract the text associated with the hyperlink
                hyperlink_text = ''.join(node.text for node in hyperlink.findall(".//w:t", namespaces=p_xml.nsmap))
                
                # Store the hyperlink text and URL in the dictionary
                hyperlinks[hyperlink_text] = url

    print("Hyperlinks extracted successfully.")
    return hyperlinks

def change_url(urls: dict[str, str], new_url_pice:str) -> dict[str, str]:
    """
    Takes a dictionary and reviews the URLs within it so that they can be changed to the correct URLs.

    :param urls: A dictionary with the text and the associated URLs. The keys are the text and the values are the URLs.
    :type urls: dict[str, str]
    :param new_url_pice: this is the new url you are wanting to use. If there is a begining and end piece it is seperated with ::
    :type new_url_pice: str
    :return: A dictionary containing hyperlink texts and their new corresponding URLs.
    :rtype: dict[str, str]
    """
    # Create a new dict for the updated URLs
    new_urls = {}

    for text, url in urls.items():
        new_url = edit_url(url,new_url_pice)

        # Add the updated URL to the new dictionary
        new_urls[text] = new_url

    return new_urls

def edit_url(old_url: str, new_url: str) -> str:
    """
    Takes the previous URL and changes it to a new link format.

    :param old_url: The old URL you want to change.
    :type old_url: str
    :param new_url: The new base URL to which the final section of the old URL will be appended.
    :type new_url: str
    :return: Returns a string with the new section of the URL attached to the desired piece.
    :rtype: str
    """
    # Step 1: Remove the '.pdf' extension if it exists at the end of the URL
    old_url = re.sub(r'\.pdf$', '', old_url)

    # Step 2: Extract the last section of the URL after the last '/'
    # This gets the last part of the URL after the final '/'
    last_section = old_url.rstrip('/').split('/')[-1]

    # Step 3: Combine the new base URL with the extracted last section
    new_full_url = new_url.rstrip('/') + '/' + last_section

    return new_full_url
    

def update_urls_in_doc(doc: Document, modified_urls: dict[str, str]) -> Document:
    """
    This function replaces the URLs within the Document with the edited ones in the modified_urls dictionary.

    :param doc: The Word document object that is opened and being edited.
    :type doc: Document
    :param modified_urls: A dictionary where the keys are the hyperlink texts and the values are the new URLs to replace the old ones.
    :type modified_urls: dict[str, str]
    :return: Returns the updated Document object.
    :rtype: Document
    """
    # Access the document's relationship dictionary
    rels = doc.part.rels

    # Iterate through each paragraph in the document
    for paragraph in doc.paragraphs:
        # Parse the paragraph XML
        p_xml = paragraph._element
        
        # Find all hyperlink tags within the paragraph
        for hyperlink in p_xml.findall(".//w:hyperlink", namespaces=p_xml.nsmap):
            # Extract the relationship ID of the hyperlink
            r_id = hyperlink.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
            
            if r_id and r_id in rels:
                # Extract the URL using the relationship ID
                current_url = rels[r_id]._target
                
                # Find the associated text with this hyperlink
                hyperlink_text = ''.join(node.text for node in hyperlink.findall(".//w:t", namespaces=p_xml.nsmap))

                # Check if the text has a corresponding new URL to replace
                if hyperlink_text in modified_urls:
                    new_url = modified_urls[hyperlink_text]  # Get the new URL
                    
                    # Update the relationship target to the new URL
                    rels[r_id]._target = new_url
    return doc

def new_path(path:str) -> str:
    """
    Changes the current path that a user inputs to create a new document

    :param path: this is the path to the current document saved as a string
    :type path: str
    :return: returns a new path
    :rtype: str

    :Example:
    >>> path = "C:\\hello.docx"
    >>> path_new = new_path(path)
    >>> assert path_new == "C:\\hello new.docx"
    """
    # Regular expression to split the path into directory, filename, and extension
    match = re.match(r"^(.*[\\/])?([^\\/]+)(\.[^\\/]+)?$", path)
    
    if match:
        # Extract directory, filename, and extension
        directory = match.group(1) or ""  # Includes trailing slash/backslash
        filename = match.group(2) or ""   # Filename without extension
        extension = match.group(3) or ""  # Extension (e.g., '.docx')

        # Construct new filename by appending ' new' before the extension
        new_filename = f"{filename} new{extension}"

        # Combine directory and new filename to form the new path
        new_path = f"{directory}{new_filename}"

        return new_path
    else:
        # If the path does not match the pattern, return the original path
        return path


def save_doc(doc: Document, path: str) -> None:
    """
    Saves the modified Word document to the specified path.

    :param doc: The Document object to be saved.
    :type doc: Document
    :param path: The file path where the document will be saved.
    :type path: str
    :return: None
    """
    try:
        doc.save(path)
        print(f"Document saved successfully at {path}.")
    except Exception as e:
        print(f"An error occurred while saving the document: {e}")
        raise

if __name__ == '__main__':
    main()
